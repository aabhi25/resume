#!/usr/bin/env python3
import sys
import json
import os
import html
from pathlib import Path
from jinja2 import Environment, FileSystemLoader, select_autoescape
import weasyprint
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class ResumeFormatter:
    def __init__(self):
        # Setup Jinja2 environment with autoescape enabled for security
        template_dir = Path(__file__).parent.parent / 'templates'
        self.env = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=select_autoescape(['html', 'xml'])
        )
        
        # Output directory
        self.output_dir = Path('outputs')
        self.output_dir.mkdir(exist_ok=True)
    
    def format_resume(self, job_data: dict, output_format: str, template_name: str) -> str:
        """Format resume using template and return file path"""
        
        if job_data.get('generationMode') == 'preserve':
            return self.preserve_format(job_data, output_format)
        else:
            return self.apply_template(job_data, output_format, template_name)
    
    def apply_template(self, job_data: dict, output_format: str, template_name: str) -> str:
        """Apply HTML template and convert to desired format"""
        
        # Parse enhanced content
        try:
            enhanced_data = json.loads(job_data['enhancedContent'])
        except:
            enhanced_data = job_data.get('parsedData', {})
        
        # Get template
        template = self.env.get_template(f'{template_name}.html')
        
        # Render HTML
        html_content = template.render(
            data=enhanced_data,
            job=job_data,
            match_score=job_data.get('matchScore', 0)
        )
        
        # Generate output file
        if output_format.lower() == 'pdf':
            return self.html_to_pdf(html_content, job_data['id'])
        elif output_format.lower() == 'docx':
            return self.html_to_docx(enhanced_data, job_data['id'])
        else:
            # Return HTML
            output_path = self.output_dir / f"resume_{job_data['id']}.html"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return str(output_path)
    
    def preserve_format(self, job_data: dict, output_format: str) -> str:
        """Preserve original format and only update content"""
        
        # Parse enhanced content with fallback to parsed data
        try:
            enhanced_data = json.loads(job_data['enhancedContent'])
        except (json.JSONDecodeError, KeyError):
            enhanced_data = job_data.get('parsedData', {})
        
        if output_format.lower() == 'docx':
            return self.create_format_preserved_docx(enhanced_data, job_data)
        else:
            return self.create_format_preserved_pdf(enhanced_data, job_data)
    
    def html_to_pdf(self, html_content: str, job_id: str) -> str:
        """Convert HTML to PDF using WeasyPrint"""
        
        output_path = self.output_dir / f"resume_{job_id}.pdf"
        
        try:
            weasyprint.HTML(string=html_content).write_pdf(str(output_path))
        except Exception as e:
            # Fallback: create simple text-based PDF
            print(f"WeasyPrint failed: {e}, creating simple PDF", file=sys.stderr)
            return self.create_fallback_pdf(html_content, job_id)
        
        return str(output_path)
    
    def html_to_docx(self, data: dict, job_id: str) -> str:
        """Convert data to DOCX format"""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Resume', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add summary
        if data.get('summary'):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(data['summary'])
        
        # Add work experience
        if data.get('workExperience'):
            doc.add_heading('Work Experience', level=1)
            for exp in data['workExperience']:
                # Company and position
                exp_para = doc.add_paragraph()
                exp_para.add_run(f"{exp.get('company', '')} - {exp.get('position', '')}").bold = True
                if exp.get('duration'):
                    exp_para.add_run(f" ({exp['duration']})")
                
                # Responsibilities
                for resp in exp.get('responsibilities', []):
                    doc.add_paragraph(resp, style='List Bullet')
        
        # Add education
        if data.get('education'):
            doc.add_heading('Education', level=1)
            for edu in data['education']:
                edu_text = f"{edu.get('degree', '')} - {edu.get('institution', '')}"
                if edu.get('year'):
                    edu_text += f" ({edu['year']})"
                doc.add_paragraph(edu_text)
        
        # Add skills
        if data.get('skills'):
            doc.add_heading('Skills', level=1)
            skills_text = ', '.join(data['skills'])
            doc.add_paragraph(skills_text)
        
        # Add projects
        if data.get('projects'):
            doc.add_heading('Projects', level=1)
            for project in data['projects']:
                proj_para = doc.add_paragraph()
                proj_para.add_run(project.get('name', '')).bold = True
                if project.get('description'):
                    doc.add_paragraph(project['description'])
                if project.get('technologies'):
                    tech_text = 'Technologies: ' + ', '.join(project['technologies'])
                    doc.add_paragraph(tech_text)
        
        # Save document
        output_path = self.output_dir / f"resume_{job_id}.docx"
        doc.save(str(output_path))
        
        return str(output_path)
    
    def create_simple_docx(self, data: dict, job_id: str) -> str:
        """Create simple DOCX preserving original format style"""
        return self.html_to_docx(data, job_id)
    
    def create_simple_pdf(self, data: dict, job_id: str) -> str:
        """Create simple PDF with basic formatting"""
        
        # Create basic HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #333; border-bottom: 2px solid #333; }}
                h2 {{ color: #666; margin-top: 20px; }}
                .company {{ font-weight: bold; }}
                .duration {{ color: #666; }}
                ul {{ margin: 5px 0; }}
            </style>
        </head>
        <body>
            <h1>Resume</h1>
            
            {f'<h2>Professional Summary</h2><p>{html.escape(data.get("summary", ""))}</p>' if data.get('summary') else ''}
            
            {self.format_work_experience_html(data.get('workExperience', []))}
            
            {self.format_education_html(data.get('education', []))}
            
            {f'<h2>Skills</h2><p>{html.escape(", ".join(data.get("skills", [])))}</p>' if data.get('skills') else ''}
            
            {self.format_projects_html(data.get('projects', []))}
        </body>
        </html>
        """
        
        return self.html_to_pdf(html_content, job_id)
    
    def create_format_preserved_pdf(self, data: dict, job_data: dict) -> str:
        """Create PDF that preserves original format structure"""
        
        # Use original resume structure with minimal formatting
        original_text = job_data.get('originalResume', '')
        
        # Create HTML that mimics the original structure
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ 
                    font-family: 'Times New Roman', serif; 
                    font-size: 12pt;
                    line-height: 1.4;
                    margin: 1in;
                    color: #000;
                }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .name {{ font-size: 18pt; font-weight: bold; margin-bottom: 5px; }}
                .contact {{ font-size: 11pt; margin: 2px 0; }}
                .section {{ margin-bottom: 15px; }}
                .section-title {{ 
                    font-weight: bold; 
                    font-size: 12pt; 
                    margin-bottom: 8px;
                    text-transform: uppercase;
                }}
                .job-title {{ font-weight: bold; margin-bottom: 2px; }}
                .company-info {{ margin-bottom: 5px; }}
                .responsibility {{ margin-left: 15px; margin-bottom: 3px; }}
                .education-item {{ margin-bottom: 5px; }}
                .skills-list {{ margin-left: 15px; }}
                p {{ margin: 3px 0; }}
            </style>
        </head>
        <body>
            {self.format_preserved_content(data, original_text, job_data)}
        </body>
        </html>
        """
        
        return self.html_to_pdf(html_content, job_data['id'])
    
    def create_format_preserved_docx(self, data: dict, job_data: dict) -> str:
        """Create DOCX that preserves original format structure"""
        
        doc = Document()
        
        # Extract user name and contact info from data
        name = self.extract_name_from_data(data, job_data)
        contact_info = self.extract_contact_from_data(data, job_data)
        
        # Add name as header (from user data)
        if name:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            name_run = name_para.add_run(name.upper())
            name_run.bold = True
            name_run.font.size = Pt(16)
        
        # Add contact info from user data
        for contact_line in contact_info:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            contact_para.add_run(contact_line)
        
        # Add summary with original style
        if data.get('summary'):
            doc.add_paragraph()  # Space
            summary_heading = doc.add_paragraph()
            summary_heading.add_run('SUMMARY').bold = True
            doc.add_paragraph(data['summary'])
        
        # Add work experience in original format
        if data.get('workExperience'):
            doc.add_paragraph()  # Space
            exp_heading = doc.add_paragraph()
            exp_heading.add_run('EXPERIENCE').bold = True
            
            for exp in data['workExperience']:
                # Job title and dates on same line
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(f"{exp.get('duration', '')} | {exp.get('position', '')}")
                job_run.bold = True
                
                # Company name
                company_para = doc.add_paragraph()
                company_para.add_run(f"{exp.get('company', '')}")
                
                # Responsibilities
                for resp in exp.get('responsibilities', []):
                    resp_para = doc.add_paragraph(resp)
                    resp_para.style = 'List Bullet'
        
        # Add education
        if data.get('education'):
            doc.add_paragraph()  # Space
            edu_heading = doc.add_paragraph()
            edu_heading.add_run('EDUCATION').bold = True
            
            for edu in data['education']:
                edu_para = doc.add_paragraph()
                edu_para.add_run(f"{edu.get('year', '')} | {edu.get('degree', '')}")
                # Add proper line break instead of \n
                edu_para.add_run().add_break()
                edu_para.add_run(f"{edu.get('institution', '')}")
        
        # Add skills
        if data.get('skills'):
            doc.add_paragraph()  # Space
            skills_heading = doc.add_paragraph()
            skills_heading.add_run('SKILLS').bold = True
            
            skills_para = doc.add_paragraph()
            skills_text = ' • '.join(data['skills'])
            skills_para.add_run(skills_text)
        
        # Add projects
        if data.get('projects'):
            doc.add_paragraph()  # Space
            proj_heading = doc.add_paragraph()
            proj_heading.add_run('PROJECTS').bold = True
            
            for project in data['projects']:
                proj_para = doc.add_paragraph()
                proj_para.add_run(project.get('name', '')).bold = True
                if project.get('description'):
                    doc.add_paragraph(project['description'])
        
        # Save document
        output_path = self.output_dir / f"resume_preserved_{job_data['id']}.docx"
        doc.save(str(output_path))
        
        return str(output_path)
    
    def format_preserved_content(self, data: dict, original_text: str, job_data: dict = None) -> str:
        """Format content to preserve original structure"""
        
        # Extract user name and contact info from actual data
        job_data = job_data or {}
        name = self.extract_name_from_data(data, job_data)
        contact_info = self.extract_contact_from_data(data, job_data)
        
        content = '<div class="header">'
        if name:
            content += f'<div class="name">{html.escape(name.upper())}</div>'
        for contact_line in contact_info:
            content += f'<div class="contact">{html.escape(contact_line)}</div>'
        content += '</div>'
        
        # Summary
        if data.get('summary'):
            content += '<div class="section">'
            content += '<div class="section-title">Summary</div>'
            content += f'<p>{html.escape(data["summary"])}</p>'
            content += '</div>'
        
        # Work Experience in original format
        if data.get('workExperience'):
            content += '<div class="section">'
            content += '<div class="section-title">Experience</div>'
            
            for exp in data['workExperience']:
                content += f'<div class="job-title">{html.escape(exp.get("duration", ""))} | {html.escape(exp.get("position", ""))}</div>'
                content += f'<div class="company-info">{html.escape(exp.get("company", ""))}</div>'
                
                for resp in exp.get('responsibilities', []):
                    content += f'<div class="responsibility">• {html.escape(resp)}</div>'
                content += '<br>'
            content += '</div>'
        
        # Education
        if data.get('education'):
            content += '<div class="section">'
            content += '<div class="section-title">Education</div>'
            
            for edu in data['education']:
                content += f'<div class="education-item">'
                content += f'{html.escape(str(edu.get("year", "")))} | {html.escape(edu.get("degree", ""))}<br>'
                content += f'{html.escape(edu.get("institution", ""))}'
                content += '</div>'
            content += '</div>'
        
        # Skills
        if data.get('skills'):
            content += '<div class="section">'
            content += '<div class="section-title">Skills</div>'
            content += f'<div class="skills-list">{html.escape(" • ".join(data["skills"]))}</div>'
            content += '</div>'
        
        # Projects
        if data.get('projects'):
            content += '<div class="section">'
            content += '<div class="section-title">Projects</div>'
            
            for project in data['projects']:
                content += f'<div class="job-title">{html.escape(project.get("name", ""))}</div>'
                if project.get('description'):
                    content += f'<p>{html.escape(project["description"])}</p>'
                if project.get('technologies'):
                    tech_text = 'Technologies: ' + ', '.join([html.escape(tech) for tech in project['technologies']])
                    content += f'<p><em>{tech_text}</em></p>'
                content += '<br>'
            content += '</div>'
        
        return content
    
    def extract_name_from_data(self, data: dict, job_data: dict) -> str:
        """Extract user name from resume data"""
        
        # Try to get name from different possible fields
        name = data.get('name') or data.get('fullName') or data.get('candidateName')
        
        # If no structured name, try to extract from original resume text
        if not name and job_data.get('originalResume'):
            original_text = job_data['originalResume']
            lines = original_text.split('\n')[:5]  # Check first 5 lines
            for line in lines:
                line = line.strip()
                # Look for name patterns (avoid emails, phone numbers, addresses)
                if (line and len(line.split()) <= 3 and 
                    not '@' in line and not any(char.isdigit() for char in line) and
                    len(line) > 5 and not line.lower().startswith(('phone', 'email', 'address'))):
                    name = line
                    break
        
        return name or "Resume"  # Fallback to generic name
    
    def extract_contact_from_data(self, data: dict, job_data: dict) -> list:
        """Extract contact information from resume data"""
        
        contact_info = []
        
        # Try to extract from structured data
        if data.get('summary'):
            summary_lines = data['summary'].split('\n')[:2]
            for line in summary_lines:
                if line.strip() and any(word in line.lower() for word in ['manager', 'engineer', 'developer', 'analyst', 'specialist']):
                    contact_info.append(line.strip())
                    break
        
        # Add location if available
        location = data.get('location') or data.get('address')
        email = data.get('email') or data.get('emailAddress')  
        phone = data.get('phone') or data.get('phoneNumber')
        
        contact_parts = []
        if location:
            contact_parts.append(location)
        if email:
            contact_parts.append(email)
        if phone:
            contact_parts.append(phone)
        
        if contact_parts:
            contact_info.append(' | '.join(contact_parts))
        
        # If no contact info found, try to extract from original resume
        if not contact_info and job_data.get('originalResume'):
            original_text = job_data['originalResume']
            lines = original_text.split('\n')[:10]
            
            for line in lines:
                line = line.strip()
                # Look for email patterns
                if '@' in line and '.' in line:
                    contact_info.append(line)
                    break
        
        return contact_info or ["Professional Resume"]  # Fallback
    
    def format_work_experience_html(self, experiences: list) -> str:
        """Format work experience as HTML"""
        if not experiences:
            return ''
        
        out = '<h2>Work Experience</h2>'
        for exp in experiences:
            out += f'<div class="experience">'
            out += f'<div class="company">{html.escape(exp.get("company", ""))} - {html.escape(exp.get("position", ""))}</div>'
            if exp.get('duration'):
                out += f'<div class="duration">{html.escape(exp["duration"])}</div>'
            
            if exp.get('responsibilities'):
                out += '<ul>'
                for resp in exp['responsibilities']:
                    out += f'<li>{html.escape(resp)}</li>'
                out += '</ul>'
            out += '</div><br>'
        
        return out
    
    def format_education_html(self, education: list) -> str:
        """Format education as HTML"""
        if not education:
            return ''
        
        out = '<h2>Education</h2>'
        for edu in education:
            edu_text = f"{html.escape(edu.get('degree', ''))} - {html.escape(edu.get('institution', ''))}"
            if edu.get('year'):
                edu_text += f" ({html.escape(str(edu['year']))})"
            out += f'<p>{edu_text}</p>'
        
        return out
    
    def format_projects_html(self, projects: list) -> str:
        """Format projects as HTML"""
        if not projects:
            return ''
        
        out = '<h2>Projects</h2>'
        for project in projects:
            out += f'<div class="project">'
            out += f'<div class="company">{html.escape(project.get("name", ""))}</div>'
            if project.get('description'):
                out += f'<p>{html.escape(project["description"])}</p>'
            if project.get('technologies'):
                escaped_techs = [html.escape(tech) for tech in project['technologies']]
                tech_text = 'Technologies: ' + ', '.join(escaped_techs)
                out += f'<p><em>{tech_text}</em></p>'
            out += '</div><br>'
        
        return out
    
    def create_fallback_pdf(self, html_content: str, job_id: str) -> str:
        """Fallback PDF creation without WeasyPrint"""
        # Create a simple text file as fallback
        output_path = self.output_dir / f"resume_{job_id}_fallback.txt"
        
        # Strip HTML tags for text version
        import re
        text_content = re.sub('<.*?>', '', html_content)
        text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        return str(output_path)

def main():
    if len(sys.argv) < 3:
        print("Usage: python resume-formatter.py <format> <template>", file=sys.stderr)
        sys.exit(1)
    
    output_format = sys.argv[1]
    template_name = sys.argv[2]
    
    try:
        # Read job data from stdin
        job_data = json.loads(sys.stdin.read())
        
        # Create formatter and generate resume
        formatter = ResumeFormatter()
        output_path = formatter.format_resume(job_data, output_format, template_name)
        
        print(output_path)
        
    except Exception as e:
        print(f"Error formatting resume: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
